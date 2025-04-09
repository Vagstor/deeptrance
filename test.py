from docx import Document
from lxml import etree
import zipfile, utils, os
from config.constants import constant
from tqdm import tqdm
import time, pyperclip

ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
}

'''document = Document()
document.save('trance.docx')'''
'''tree = ET.parse('document.xml')
root = tree.getroot()
for text in root.findall('w:t'):
    print(text.text)'''

'''def parse_docx_xml(docx_path):

    with zipfile.ZipFile(docx_path) as z:
        with z.open('word/document.xml') as f:
            xml_content = f.read()

    parser = etree.XMLParser(remove_blank_text=True)
    return etree.fromstring(xml_content, parser)

root = parse_docx_xml('orig/cnc8_070 ошибки.docx')

for paragraph in root.xpath('//w:p', namespaces=ns):
    for run in paragraph.xpath('.//w:r', namespaces=ns):
        text_nodes = run.xpath('.//w:t', namespaces=ns)
        for t in text_nodes:
            if t.text:
                print(t.text)'''
# try:
#     doc = Document(os.path.join(constant.root_dir, 'orig', 'Forklift_EN.docx'))
# except Exception as e:
#     print('Ошибка: ', e)

# client = utils.create_ds_client()

# for paragraph in tqdm(doc.paragraphs):
#     if 'Forklift size' in paragraph.text:
#         paragraph.text = paragraph.text.replace('Forklift size', 'Габариты тележки')
#     time.sleep(2)
# doc.save(os.path.join(constant.root_dir, 'translation', 'Forklift_ru'))

print(pyperclip.paste())
# for paragraph in tqdm(doc.paragraphs):
#     completion = utils.send_request(client, utils.create_message(paragraph.text))
#     with open('prompt_stream.txt', 'a', encoding='utf-8') as file:
#       file.write(completion.choices[0].message.content + '\n')
#     paragraph.text = completion.choices[0].message.content
# doc.save(os.path.join(constant.root_dir + 'translation', 'Forklift_ru'))